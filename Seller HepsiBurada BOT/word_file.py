from docx import Document
from docx.shared import Inches

document = Document()

p = document.add_paragraph()
r = p.add_run()

r.add_tex()
